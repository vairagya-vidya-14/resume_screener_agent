import os
import docx
from fpdf import FPDF

def create_txt_resume(filepath, name, title, edu, exp_yrs, skills, summary, details):
    linkedin = f"https://www.linkedin.com/in/{name.lower().replace(' ', '-')}"
    github = f"https://github.com/{name.lower().replace(' ', '')}"
    content = f"""{name}
{title} | Email: {name.lower().replace(' ', '.')}@email.com | Phone: +1-555-019-2834
LinkedIn: {linkedin} | GitHub: {github}

SUMMARY
{summary}

EDUCATION
{edu}

EXPERIENCE ({exp_yrs} Years)
{details}

SKILLS
{', '.join(skills)}
"""
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)

def create_docx_resume(filepath, name, title, edu, exp_yrs, skills, summary, details):
    doc = docx.Document()
    doc.add_heading(name, level=0)
    linkedin = f"https://www.linkedin.com/in/{name.lower().replace(' ', '-')}"
    github = f"https://github.com/{name.lower().replace(' ', '')}"
    
    doc.add_paragraph(f"{title} | Email: {name.lower().replace(' ', '.')}@email.com | Phone: +1-555-019-2834\nLinkedIn: {linkedin} | GitHub: {github}")
    
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(summary)
    
    doc.add_heading("Education", level=1)
    doc.add_paragraph(edu)
    
    doc.add_heading(f"Work Experience ({exp_yrs} Years)", level=1)
    doc.add_paragraph(details)
    
    doc.add_heading("Technical & Core Skills", level=1)
    doc.add_paragraph(", ".join(skills))
    
    doc.save(filepath)

def create_pdf_resume(filepath, name, title, edu, exp_yrs, skills, summary, details, email=None, phone=None):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, name, new_x="LMARGIN", new_y="NEXT", align='C')
    pdf.set_font("Arial", 'I', 11)
    
    cand_email = email or f"{name.lower().replace(' ', '.')}@email.com"
    cand_phone = phone or "+1-555-019-2834"
    linkedin = f"https://www.linkedin.com/in/{name.lower().replace(' ', '-')}"
    github = f"https://github.com/{name.lower().replace(' ', '')}"
    
    pdf.cell(0, 8, f"{title} | {cand_email} | {cand_phone}", new_x="LMARGIN", new_y="NEXT", align='C')
    pdf.cell(0, 8, f"LinkedIn | GitHub", new_x="LMARGIN", new_y="NEXT", align='C', link=linkedin)
    pdf.ln(5)

    pdf.set_font("Arial", 'B', 12)
    pdf.cell(0, 8, "PROFESSIONAL SUMMARY", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Arial", '', 10)
    pdf.multi_cell(0, 6, summary)
    pdf.ln(3)

    pdf.set_font("Arial", 'B', 12)
    pdf.cell(0, 8, "EDUCATION", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Arial", '', 10)
    pdf.multi_cell(0, 6, edu)
    pdf.ln(3)

    pdf.set_font("Arial", 'B', 12)
    pdf.cell(0, 8, f"WORK EXPERIENCE ({exp_yrs} YEARS)", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Arial", '', 10)
    pdf.multi_cell(0, 6, details)
    pdf.ln(3)

    pdf.set_font("Arial", 'B', 12)
    pdf.cell(0, 8, "TECHNICAL SKILLS", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Arial", '', 10)
    pdf.multi_cell(0, 6, ", ".join(skills))

    pdf.output(filepath)

def generate_all():
    target_dir = os.path.join("data", "sample_resumes")
    os.makedirs(target_dir, exist_ok=True)

    candidates = [
        {
            "filename": "candidate_01_alex_chen.pdf",
            "type": "pdf",
            "name": "Alex Chen",
            "title": "Senior AI & ML Research Engineer",
            "edu": "Master of Science in Computer Science, Stanford University (2018-2020)",
            "exp_yrs": 6,
            "skills": ["Python", "PyTorch", "TensorFlow", "Transformers", "LLM", "RAG", "LangChain", "FastAPI", "Docker", "Kubernetes", "AWS", "MLOps", "SQL", "Git"],
            "summary": "Senior AI Engineer with 6 years of experience building scalable Generative AI systems, RAG pipelines, and transformer model deployment.",
            "details": "Lead AI Engineer at TechCorp (2020-Present): Designed enterprise RAG pipeline serving 500k daily queries with PyTorch, FastAPI, and Qdrant. Fine-tuned Llama models using PEFT/LoRA."
        },
        {
            "filename": "candidate_02_sarah_jenkins.docx",
            "type": "docx",
            "name": "Sarah Jenkins",
            "title": "AI Research Scientist & NLP Specialist",
            "edu": "PhD in Computer Science (NLP Focus), MIT (2015-2020)",
            "exp_yrs": 5,
            "skills": ["Python", "PyTorch", "NLP", "Transformers", "Hugging Face", "BERT", "GPT", "Prompt Engineering", "Fine-tuning", "Scikit-Learn", "Research", "C++"],
            "summary": "PhD AI Scientist specialized in Natural Language Processing, Transformer architectures, and LLM alignment research.",
            "details": "Research Scientist at OpenAI Labs (2020-Present): Published 4 top-tier NLP papers. Optimized transformer latency by 35% using TensorRT and PyTorch."
        },
        {
            "filename": "candidate_03_marcus_vance.pdf",
            "type": "pdf",
            "name": "Marcus Vance",
            "title": "Senior Data Engineer",
            "edu": "Bachelor of Science in Information Systems, UC Berkeley (2016-2020)",
            "exp_yrs": 5,
            "skills": ["Python", "SQL", "PostgreSQL", "MongoDB", "Apache Spark", "Kafka", "Snowflake", "Docker", "Airflow", "ETL", "AWS"],
            "summary": "Data Engineer with 5 years experience designing distributed ETL pipelines and data warehouses.",
            "details": "Senior Data Engineer at DataScale (2020-Present): Built Spark pipelines processing 2TB daily log streams into Snowflake warehouse."
        },
        {
            "filename": "candidate_04_priya_sharma.txt",
            "type": "txt",
            "name": "Priya Sharma",
            "title": "Machine Learning & MLOps Engineer",
            "edu": "Master of Engineering in Artificial Intelligence, Carnegie Mellon University (2019-2021)",
            "exp_yrs": 5,
            "skills": ["Python", "PyTorch", "Scikit-Learn", "MLOps", "Docker", "Kubernetes", "MLflow", "FastAPI", "AWS", "LLM", "RAG", "Git", "CI/CD"],
            "summary": "MLOps Engineer passionate about continuous integration and automated deployment of deep learning models.",
            "details": "MLOps Lead at CloudAI (2021-Present): Deployed Kubernetes clusters for PyTorch LLM inference microservices on AWS EKS."
        },
        {
            "filename": "candidate_05_david_kim.docx",
            "type": "docx",
            "name": "David Kim",
            "title": "Backend Microservices Developer",
            "edu": "Bachelor of Science in Computer Science, UT Austin (2017-2021)",
            "exp_yrs": 4,
            "skills": ["Java", "Python", "REST API", "Microservices", "Docker", "Kubernetes", "PostgreSQL", "Redis", "Git", "Linux"],
            "summary": "Backend Software Developer skilled in Java Spring Boot and Python FastAPI microservices.",
            "details": "Software Engineer at Enterprise Software Co (2021-Present): Maintained REST APIs and database schema design."
        },
        {
            "filename": "candidate_06_elena_rodriguez.pdf",
            "type": "pdf",
            "name": "Elena Rodriguez",
            "title": "Data Analyst & Business Intelligence Specialist",
            "edu": "Bachelor of Arts in Economics & Statistics, UCLA (2018-2022)",
            "exp_yrs": 3,
            "skills": ["Python", "Pandas", "NumPy", "SQL", "Tableau", "Matplotlib", "Seaborn", "Data Analysis", "A/B Testing"],
            "summary": "Data Analyst with 3 years of experience in statistical analysis, business dashboards, and A/B testing.",
            "details": "Data Analyst at RetailInsights (2022-Present): Analyzed customer retention metrics using SQL and Python Pandas."
        },
        {
            "filename": "candidate_07_james_wilson.txt",
            "type": "txt",
            "name": "James Wilson",
            "title": "Frontend UI/UX React Developer",
            "edu": "Bachelor of Science in Web Design, Oregon State (2019-2023)",
            "exp_yrs": 2,
            "skills": ["JavaScript", "TypeScript", "React", "HTML/CSS", "Node.js", "Git", "Figma", "Web Development"],
            "summary": "Frontend Engineer focused on building reactive user interfaces in React and TypeScript.",
            "details": "Frontend Developer at WebStudio (2023-Present): Created user dashboard components."
        },
        {
            "filename": "candidate_08_amira_al_mansoor.pdf",
            "type": "pdf",
            "name": "Amira Al-Mansoor",
            "title": "Lead NLP & LLM Engineer",
            "edu": "Master of Science in Artificial Intelligence, Oxford University (2017-2019)",
            "exp_yrs": 7,
            "skills": ["Python", "PyTorch", "NLP", "Transformers", "LLM", "RAG", "Fine-tuning", "LlamaIndex", "Hugging Face", "Scikit-Learn", "Docker", "AWS", "Research"],
            "summary": "Lead AI Engineer with 7 years of deep specialization in NLP models, conversational AI, and vector retrieval.",
            "details": "Lead NLP Architect at Global AI Labs (2019-Present): Architected multi-lingual LLM RAG platform for legal compliance."
        },
        {
            "filename": "candidate_09_robert_taylor.docx",
            "type": "docx",
            "name": "Robert Taylor",
            "title": "IT Project Manager & Agile Scrum Master",
            "edu": "Bachelor of Business Administration, University of Washington (2012-2016)",
            "exp_yrs": 9,
            "skills": ["Project Management", "Agile", "Scrum", "Leadership", "Product Strategy", "Communication", "Cross-Functional Leadership"],
            "summary": "Agile Project Manager overseeing software delivery sprints and engineering team roadmaps.",
            "details": "Senior Project Manager at AgileTech (2017-Present): Managed engineering sprints across 4 cross-functional teams."
        },
        {
            "filename": "candidate_10_lisa_wang.pdf",
            "type": "pdf",
            "name": "Lisa Wang",
            "title": "Computer Vision & Deep Learning Engineer",
            "edu": "Master of Science in Robotics, Georgia Tech (2018-2020)",
            "exp_yrs": 5,
            "skills": ["Python", "PyTorch", "Computer Vision", "OpenCV", "Deep Learning", "TensorFlow", "C++", "Docker", "Linux"],
            "summary": "Computer Vision Engineer experienced in object detection, image segmentation, and edge device deployment.",
            "details": "Vision Engineer at RoboTech (2020-Present): Developed OpenCV and PyTorch YOLO models for autonomous navigation."
        },
        {
            "filename": "candidate_11_devon_brooks.txt",
            "type": "txt",
            "name": "Devon Brooks",
            "title": "Technical Content Writer & Documentation Manager",
            "edu": "Bachelor of Arts in English Literature, NYU (2018-2022)",
            "exp_yrs": 3,
            "skills": ["Technical Writing", "Communication", "Research", "Git", "Markdown", "Content Strategy"],
            "summary": "Technical Content Writer crafting clear API documentation, developer guides, and tech articles.",
            "details": "Technical Writer at DevDocs Inc (2022-Present): Authored developer documentation and API references."
        },
        {
            "filename": "candidate_12_ananya_patel.pdf",
            "type": "pdf",
            "name": "Ananya Patel",
            "title": "Senior MLOps & Generative AI Infrastructure Engineer",
            "edu": "Master of Science in Computer Engineering, University of Illinois (2017-2019)",
            "exp_yrs": 6,
            "skills": ["Python", "PyTorch", "LLM", "RAG", "MLOps", "Docker", "Kubernetes", "AWS", "Terraform", "FastAPI", "SQL", "Git", "CI/CD"],
            "summary": "AI Infrastructure Engineer with 6 years experience scaling Generative AI models and MLOps platforms in production.",
            "details": "Senior Infrastructure Lead at ScaleAI (2019-Present): Orchestrated EKS Kubernetes cluster serving LLM inference with 99.99% uptime."
        },
        {
            "filename": "candidate_13_vairagya_vidya.pdf",
            "type": "pdf",
            "name": "Vairagya Vidya",
            "title": "Python & Full Stack Developer",
            "email": "vidyavairagya@gmail.com",
            "phone": "6300957267",
            "edu": "B.Tech in Computer Science and Engineering, JNTUH College of Engineering Sultanpur (2022-2026)",
            "exp_yrs": 0,
            "skills": ["Python", "Flask", "MySQL", "HTML5", "CSS3", "JavaScript", "SQL", "Git", "GitHub", "REST API", "Java", "C"],
            "summary": "Motivated Computer Science undergraduate with hands-on experience designing, coding, and enhancing custom software solutions using Python, Flask, MySQL, and web development practices.",
            "details": "Academic & Personal Projects: StegoSafe (Steganography web app) and News Vista (Personalized News platform) using Python, Flask, MySQL, and REST APIs."
        }
    ]

    for c in candidates:
        filepath = os.path.join(target_dir, c["filename"])
        if c["type"] == "txt":
            create_txt_resume(filepath, c["name"], c["title"], c["edu"], c["exp_yrs"], c["skills"], c["summary"], c["details"])
        elif c["type"] == "docx":
            create_docx_resume(filepath, c["name"], c["title"], c["edu"], c["exp_yrs"], c["skills"], c["summary"], c["details"])
        elif c["type"] == "pdf":
            create_pdf_resume(filepath, c["name"], c["title"], c["edu"], c["exp_yrs"], c["skills"], c["summary"], c["details"], email=c.get("email"), phone=c.get("phone"))

    print(f"Successfully generated {len(candidates)} sample resumes!")

if __name__ == "__main__":
    generate_all()
