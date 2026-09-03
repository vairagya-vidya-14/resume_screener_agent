import os
import re
from typing import Dict, Any, Optional

class ResumeParser:
    """
    Unified Resume Parser supporting PDF, DOCX, and TXT format parsing with text normalization
    and hyperlinked URL extraction from annotations and XML relationships.
    """

    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt'}

    @classmethod
    def parse_file(cls, file_path: str) -> Dict[str, Any]:
        """
        Parses a single resume file and returns normalized text along with file metadata.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        filename = os.path.basename(file_path)
        ext = os.path.splitext(filename)[1].lower()

        if ext not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file format '{ext}'. Supported formats: {cls.SUPPORTED_EXTENSIONS}")

        raw_text = ""
        error_msg = None

        try:
            if ext == '.pdf':
                raw_text = cls._extract_pdf_text(file_path)
            elif ext == '.docx':
                raw_text = cls._extract_docx_text(file_path)
            elif ext == '.txt':
                raw_text = cls._extract_txt_text(file_path)
        except Exception as e:
            error_msg = str(e)
            raw_text = ""

        clean_text = cls.clean_text(raw_text)

        return {
            "file_path": os.path.abspath(file_path),
            "filename": filename,
            "format": ext[1:].upper(),
            "file_size_bytes": os.path.getsize(file_path),
            "raw_text": raw_text,
            "clean_text": clean_text,
            "word_count": len(clean_text.split()),
            "character_count": len(clean_text),
            "parse_error": error_msg
        }

    @staticmethod
    def _extract_pdf_text(file_path: str) -> str:
        """Extract text from PDF file using pypdf, including all hyperlinked annotations and URIs."""
        import pypdf
        text_parts = []
        with open(file_path, 'rb') as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text_parts.append(extracted)

                # Comprehensive annotation link extraction (/Annots -> /Link -> /URI)
                if "/Annots" in page:
                    try:
                        annots = page["/Annots"]
                        for annot in annots:
                            obj = annot.get_object() if hasattr(annot, "get_object") else annot
                            if isinstance(obj, dict):
                                # Direct /URI
                                if "/URI" in obj:
                                    text_parts.append(str(obj["/URI"]))
                                # Indirect or action /A -> /URI
                                action = obj.get("/A")
                                if action:
                                    act_obj = action.get_object() if hasattr(action, "get_object") else action
                                    if isinstance(act_obj, dict) and "/URI" in act_obj:
                                        text_parts.append(str(act_obj["/URI"]))
                    except Exception:
                        pass
        return "\n".join(text_parts)

    @staticmethod
    def _extract_docx_text(file_path: str) -> str:
        """Extract text from DOCX file using python-docx, including relationship hyperlinks."""
        import docx
        doc = docx.Document(file_path)
        text_parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text.strip())
            try:
                for rel_id in p._element.xpath('.//w:hyperlink/@r:id'):
                    if rel_id in doc.part.rels:
                        url = doc.part.rels[rel_id].target_ref
                        text_parts.append(url)
            except Exception:
                pass

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n".join(text_parts)

    @staticmethod
    def _extract_txt_text(file_path: str) -> str:
        """Extract text from TXT file."""
        encodings = ['utf-8', 'latin-1', 'cp1252']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    @staticmethod
    def clean_text(text: str) -> str:
        """Normalizes whitespace, removes control characters, and cleans text formatting."""
        if not text:
            return ""
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n\s*\n+', '\n\n', text)
        return text.strip()
